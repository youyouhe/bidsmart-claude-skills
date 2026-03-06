#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Word文档后验证工具
验证生成的Word文档是否完整、格式正确
"""

import os
import sys
from pathlib import Path


def verify_docx(docx_path):
    """验证Word文档"""
    try:
        from docx import Document
    except ImportError:
        print("错误: 未安装python-docx库")
        print("请运行: pip install python-docx")
        sys.exit(1)

    if not os.path.exists(docx_path):
        print(f"错误: 文件不存在: {docx_path}")
        sys.exit(1)

    try:
        doc = Document(docx_path)
        issues = []
        warnings = []

        # 1. 统计基本信息
        total_paragraphs = len(doc.paragraphs)
        total_tables = len(doc.tables)

        # 2. 统计标题
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        heading_count = len(headings)

        # 3. 统计图片
        images = []
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                images.append(rel.target_ref)
        image_count = len(images)

        # 4. 检查空段落比例
        empty_paragraphs = [p for p in doc.paragraphs if not p.text.strip()]
        empty_ratio = len(empty_paragraphs) / total_paragraphs if total_paragraphs > 0 else 0

        if empty_ratio > 0.3:
            warnings.append({
                'type': 'empty_paragraphs',
                'message': f'空段落比例过高: {len(empty_paragraphs)}/{total_paragraphs} ({empty_ratio*100:.1f}%)',
                'suggestion': '检查Markdown转换是否正确'
            })

        # 5. 检查标题层级
        prev_level = 0
        for i, heading in enumerate(headings):
            level = int(heading.style.name.replace('Heading ', ''))
            if level > prev_level + 1:
                warnings.append({
                    'type': 'heading_skip',
                    'message': f'标题跳级: Heading {prev_level} → Heading {level}',
                    'content': heading.text[:50],
                    'suggestion': '检查Markdown源文件标题层级'
                })
            prev_level = level

        # 6. 检查是否有文本内容
        has_content = any(p.text.strip() for p in doc.paragraphs)
        if not has_content:
            issues.append({
                'type': 'no_content',
                'message': '文档为空，没有任何文本内容',
                'suggestion': '检查Markdown源文件是否存在'
            })

        # 7. 检查文件大小
        file_size = os.path.getsize(docx_path)
        if file_size < 10 * 1024:  # 小于10KB
            warnings.append({
                'type': 'small_file',
                'message': f'文件大小异常: {file_size} bytes',
                'suggestion': '可能转换不完整，检查源文件'
            })

        return {
            'file': os.path.basename(docx_path),
            'stats': {
                'total_paragraphs': total_paragraphs,
                'heading_count': heading_count,
                'table_count': total_tables,
                'image_count': image_count,
                'empty_paragraphs': len(empty_paragraphs),
                'file_size': file_size,
                'file_size_mb': file_size / (1024 * 1024)
            },
            'issues': issues,
            'warnings': warnings,
            'success': True
        }

    except Exception as e:
        return {
            'file': os.path.basename(docx_path),
            'error': str(e),
            'success': False
        }


def main():
    if len(sys.argv) < 2:
        print("用法: python verify_docx.py <Word文档路径>")
        print("示例: python verify_docx.py 响应文件/响应文件-XX公司-XX项目.docx")
        sys.exit(1)

    docx_path = sys.argv[1]

    print("=== Word文档验证工具 ===")
    print(f"文件: {docx_path}\n")

    result = verify_docx(docx_path)

    if not result['success']:
        print(f"✗ 验证失败: {result['error']}")
        sys.exit(1)

    print("=" * 60)
    print("文档统计:")
    stats = result['stats']
    print(f"  段落总数: {stats['total_paragraphs']}")
    print(f"  标题数量: {stats['heading_count']}")
    print(f"  表格数量: {stats['table_count']}")
    print(f"  图片数量: {stats['image_count']}")
    print(f"  空段落数: {stats['empty_paragraphs']}")
    print(f"  文件大小: {stats['file_size_mb']:.2f} MB")

    if result['issues']:
        print("\n" + "=" * 60)
        print(f"✗ 发现 {len(result['issues'])} 个问题:")
        for issue in result['issues']:
            print(f"\n  类型: {issue['type']}")
            print(f"  问题: {issue['message']}")
            if issue.get('suggestion'):
                print(f"  建议: {issue['suggestion']}")

    if result['warnings']:
        print("\n" + "=" * 60)
        print(f"⚠️  发现 {len(result['warnings'])} 个警告:")
        for warning in result['warnings']:
            print(f"\n  类型: {warning['type']}")
            print(f"  警告: {warning['message']}")
            if warning.get('suggestion'):
                print(f"  建议: {warning['suggestion']}")

    if not result['issues'] and not result['warnings']:
        print("\n" + "=" * 60)
        print("✓ 文档验证通过，未发现问题")

    print("=" * 60)

    # 如果有ERROR级别问题，返回非0退出码
    if result['issues']:
        sys.exit(1)
    else:
        sys.exit(0)


if __name__ == "__main__":
    main()
