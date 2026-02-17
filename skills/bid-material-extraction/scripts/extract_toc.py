"""
Step 1: Extract table of contents from a Word (.docx) file.

Builds a structured TOC by scanning Heading styles (primary) or text patterns (fallback).
Generates sequential numbering matching the document hierarchy.

Usage:
    python extract_toc.py <docx_path> [--output toc.json]
"""
import argparse
import json
import re

from docx import Document

CHINESE_NUMS = {1: '一', 2: '二', 3: '三', 4: '四', 5: '五',
                6: '六', 7: '七', 8: '八', 9: '九', 10: '十'}


def extract_toc(docx_path):
    """Extract TOC from a Word document."""
    doc = Document(docx_path)

    # ── Strategy 1: Build TOC from Heading styles ──
    entries = _from_heading_styles(doc)
    if entries:
        return entries

    # ── Strategy 2: Parse explicit TOC section ──
    entries = _from_toc_section(doc)
    if entries:
        return entries

    # ── Strategy 3: Scan body for numbered section headers ──
    return _from_numbered_headers(doc)


def _from_heading_styles(doc):
    """Build TOC from paragraphs with Heading 1-6 styles."""
    entries = []
    counters = [0, 0, 0, 0, 0, 0]  # h1..h6
    current_part = None

    for pi, para in enumerate(doc.paragraphs):
        style = para.style.name if para.style else ''
        if not style.startswith('Heading '):
            continue

        text = para.text.strip()
        if not text:
            continue

        try:
            level = int(style.split()[-1])
        except (ValueError, IndexError):
            continue

        if level < 1 or level > 6:
            continue

        # Increment counter for this level, reset all deeper levels
        counters[level - 1] += 1
        for i in range(level, 6):
            counters[i] = 0

        # Build number string
        if level == 1:
            num = CHINESE_NUMS.get(counters[0], str(counters[0]))
            number = f"{num}、"
            current_part = text
        else:
            # Level 2+ uses dotted arabic notation based on parent counters
            # h2 → "N", h3 → "N.M", h4 → "N.M.K"
            parts = [str(counters[i]) for i in range(1, level)]
            number = '.'.join(parts)

        entries.append({
            'number': number,
            'title': text,
            'page': None,
            'level': level,
            'part': current_part,
            'para_index': pi,
        })

    return entries


def _from_toc_section(doc):
    """Parse a formal TOC section (目录 heading followed by TOC entries)."""
    entries = []
    in_toc = False
    current_part = None

    for para in doc.paragraphs:
        text = para.text.strip()
        style = para.style.name.lower() if para.style else ''

        # Detect TOC by style
        if 'toc' in style and text:
            entry = _parse_toc_line(text)
            if entry:
                if entry['level'] == 1:
                    current_part = entry['title']
                entry['part'] = current_part
                entries.append(entry)
            continue

        # Detect TOC by heading
        if text in ('目录', '目 录') and not in_toc:
            in_toc = True
            continue

        if not in_toc:
            continue
        if not text:
            continue

        norm = text.replace('\t', '  ')
        norm = re.sub(r'\.{3,}', '  ', norm)
        norm = re.sub(r'…+', '  ', norm)

        entry = _parse_toc_line(norm)
        if entry:
            if entry['level'] == 1:
                current_part = entry['title']
            entry['part'] = current_part
            entries.append(entry)
        elif entries and len(entries) > 3:
            break

    return entries


def _parse_toc_line(text):
    """Try to parse a single TOC line."""
    text = text.strip()

    # Chinese major section: 一、标题  页码
    m = re.match(r'^([一二三四五六七八九十]+)、\s*(.+)', text)
    if m:
        title, page = _extract_page(m.group(2))
        if title:
            return {'number': m.group(1) + '、', 'title': title,
                    'page': page, 'level': 1}

    # Arabic: 10.1.  标题  页码
    m = re.match(r'^(\d+(?:\.\d+)*)[\.\．]\s*(.+)', text)
    if m:
        num = m.group(1)
        title, page = _extract_page(m.group(2))
        if title and int(num.split('.')[0]) < 100:
            return {'number': num, 'title': title, 'page': page,
                    'level': len(num.split('.')) + 1}
    return None


def _extract_page(rest):
    """Extract page number from end of string."""
    rest = rest.rstrip()
    m = re.search(r'\s{2,}(\d{1,4})\s*$', rest)
    if m:
        return rest[:m.start()].strip(), int(m.group(1))
    m = re.search(r'\s(\d{1,4})$', rest)
    if m and int(m.group(1)) < 2000:
        return rest[:m.start()].strip(), int(m.group(1))
    return rest.strip(), None


def _from_numbered_headers(doc):
    """Fallback: scan body paragraphs for numbered section headers."""
    entries = []
    current_part = None

    for pi, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text or len(text) > 100:
            continue

        m = re.match(r'^([一二三四五六七八九十]+)、\s*(.+)', text)
        if m:
            title = m.group(2).strip()
            current_part = title
            entries.append({
                'number': m.group(1) + '、', 'title': title,
                'page': None, 'level': 1, 'part': current_part,
                'para_index': pi,
            })
            continue

        m = re.match(r'^(\d+(?:\.\d+)*)[\.\．]\s*(.+)', text)
        if m:
            num = m.group(1)
            title = m.group(2).strip()
            top = int(num.split('.')[0])
            if top < 100 and len(title) < 80:
                entries.append({
                    'number': num, 'title': title,
                    'page': None, 'level': len(num.split('.')) + 1,
                    'part': current_part, 'para_index': pi,
                })

    return entries


def main():
    parser = argparse.ArgumentParser(description="Extract TOC from Word document")
    parser.add_argument('docx_path')
    parser.add_argument('--output', '-o', default='toc.json')
    args = parser.parse_args()

    print(f"Extracting TOC from {args.docx_path}...")
    entries = extract_toc(args.docx_path)

    output = {
        'source': args.docx_path,
        'entry_count': len(entries),
        'entries': entries,
    }

    with open(args.output, 'w', encoding='utf-8') as f:
        json.dump(output, f, ensure_ascii=False, indent=2)

    print(f"Found {len(entries)} entries → {args.output}")
    for e in entries:
        indent = '  ' * (e['level'] - 1)
        page = f" (p{e['page']})" if e.get('page') else ''
        print(f"  {indent}{e['number']}  {e['title'][:60]}{page}")


if __name__ == '__main__':
    main()
