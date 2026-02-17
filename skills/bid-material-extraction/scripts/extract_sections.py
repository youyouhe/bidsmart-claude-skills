"""
Step 3: Extract sections from a Word document based on an extraction plan.

Uses toc.json (from Step 1) for precise section positioning, then extracts
text and/or images from each planned section as individual files.

Usage:
    python extract_sections.py <docx_path> --plan extraction_plan.json --toc toc.json [--output-dir data]
"""
import argparse
import json
import os
import re

from docx import Document
from docx.oxml.ns import qn


MIN_IMAGE_BYTES = 5000
ALLOWED_EXTS = {'png', 'jpg', 'jpeg', 'gif', 'bmp', 'tiff', 'tif'}


# ── Helpers ─────────────────────────────────────────────────────────────

def get_para_text(elem):
    """Get plain text from a w:p XML element."""
    texts = []
    for r in elem.iter(qn('w:t')):
        if r.text:
            texts.append(r.text)
    return ''.join(texts)


def format_table(table):
    """Format a python-docx Table as pipe-separated text."""
    lines = []
    for row in table.rows:
        seen = set()
        cells = []
        for cell in row.cells:
            cid = id(cell._tc)
            if cid not in seen:
                seen.add(cid)
                cells.append(cell.text.strip().replace('\n', ' '))
        lines.append(' | '.join(cells))
    return '\n'.join(lines)


def safe_filename(s):
    """Make string safe for filenames."""
    s = re.sub(r'[\\/:*?"<>|\u201c\u201d\u2018\u2019]', '_', s)
    s = re.sub(r'[（）()\[\]【】]', '_', s)
    s = re.sub(r'\s+', '_', s)
    s = s.strip('_.')
    return s[:80]


# ── Position mapping ────────────────────────────────────────────────────

def build_para_to_elem_map(doc):
    """Map paragraph index (doc.paragraphs[i]) → body element index."""
    body = doc.element.body
    elements = list(body)
    para_to_elem = {}
    pi = 0
    for ei, elem in enumerate(elements):
        if elem.tag == qn('w:p'):
            para_to_elem[pi] = ei
            pi += 1
    return elements, para_to_elem


def _is_descendant(child_num, parent_num):
    """Check if child_num is a sub-section of parent_num.

    e.g. '13.1.1' is descendant of '13.1' and '13',
         '13.2' is NOT a descendant of '13.1'.
    """
    # Chinese numbered sections (一、) have no children with same prefix
    if '、' in parent_num:
        return False
    return child_num.startswith(parent_num + '.')


def resolve_positions(plan_entries, toc_entries, para_to_elem, total_elements):
    """Map each plan entry to (start_elem_idx, end_elem_idx) using TOC data."""
    # Build lookup: (number, title) → para_index from toc
    toc_lookup = {}
    # sorted list of (para_index, number) for all toc entries
    toc_by_idx = []
    for te in toc_entries:
        key = (te['number'], te['title'])
        pi = te.get('para_index')
        if pi is not None:
            toc_lookup[key] = pi
            toc_by_idx.append((pi, te['number']))
    toc_by_idx.sort()

    # For each plan entry, find its start position and the next section's start
    positions = []
    for entry in plan_entries:
        key = (entry['number'], entry['title'])
        pi = toc_lookup.get(key)
        if pi is None:
            print(f"  ⚠ Section {entry['number']} '{entry['title']}' not in toc.json")
            continue

        start_ei = para_to_elem.get(pi)
        if start_ei is None:
            print(f"  ⚠ Section {entry['number']} para_index {pi} not mapped")
            continue

        # End = next non-descendant toc entry's element index.
        # This ensures requesting "13.1" captures all sub-sections (13.1.1, 13.1.2, ...)
        # and stops at the next sibling or higher-level entry (13.2).
        next_pi = None
        for tp, tn in toc_by_idx:
            if tp > pi and not _is_descendant(tn, entry['number']):
                next_pi = tp
                break

        if next_pi is not None and next_pi in para_to_elem:
            end_ei = para_to_elem[next_pi]
        else:
            end_ei = total_elements

        positions.append((start_ei, end_ei, entry))

    return positions


# ── Content extraction ──────────────────────────────────────────────────

def extract_text_from_range(doc, elements, start, end):
    """Extract readable text from body elements in [start, end)."""
    parts = []
    table_map = {id(t._element): t for t in doc.tables}

    for idx in range(start, end):
        elem = elements[idx]
        if elem.tag == qn('w:p'):
            text = get_para_text(elem).strip()
            if text:
                parts.append(text)
        elif elem.tag == qn('w:tbl'):
            t = table_map.get(id(elem))
            if t:
                parts.append(format_table(t))

    return '\n'.join(parts)


def extract_images_from_range(doc, elements, start, end):
    """Extract images from body elements in [start, end)."""
    images = []
    seen = set()

    for idx in range(start, end):
        elem = elements[idx]
        for blip in elem.iter(qn('a:blip')):
            rId = blip.get(qn('r:embed'))
            if not rId or rId in seen:
                continue
            if rId not in doc.part.rels:
                continue
            seen.add(rId)

            try:
                part = doc.part.rels[rId].target_part
                data = part.blob
                ext = part.partname.split('.')[-1].lower()

                if ext not in ALLOWED_EXTS:
                    continue
                if len(data) < MIN_IMAGE_BYTES:
                    continue

                images.append((data, ext))
            except Exception:
                pass

    return images


# ── Main ────────────────────────────────────────────────────────────────

def extract_by_plan(docx_path, plan_path, toc_path, output_dir):
    os.makedirs(output_dir, exist_ok=True)

    with open(plan_path, 'r', encoding='utf-8') as f:
        plan = json.load(f)
    with open(toc_path, 'r', encoding='utf-8') as f:
        toc = json.load(f)

    extractions = plan.get('extractions', [])
    toc_entries = toc.get('entries', [])
    if not extractions:
        print("Empty extraction plan.")
        return

    print(f"Loading {docx_path}...")
    doc = Document(docx_path)

    print("Building position map...")
    elements, para_to_elem = build_para_to_elem_map(doc)

    print(f"Resolving {len(extractions)} sections...")
    positions = resolve_positions(extractions, toc_entries, para_to_elem, len(elements))
    print(f"  Resolved {len(positions)} of {len(extractions)} sections\n")

    results = []

    for start_ei, end_ei, entry in positions:
        num = entry['number']
        title = entry['title']
        name = safe_filename(entry.get('output_name', title))
        prefix = f"{num}-{name}" if '、' not in num else name
        want_text = entry.get('extract_text', True)
        want_images = entry.get('extract_images', True)

        print(f"  {num}. {title[:50]}")
        files = []

        if want_text:
            text = extract_text_from_range(doc, elements, start_ei, end_ei)
            if text.strip():
                fname = f"{prefix}.txt"
                path = os.path.join(output_dir, fname)
                with open(path, 'w', encoding='utf-8') as f:
                    f.write(text)
                files.append(fname)
                print(f"    → {fname} ({len(text):,} chars)")

        if want_images:
            imgs = extract_images_from_range(doc, elements, start_ei, end_ei)
            for ii, (data, ext) in enumerate(imgs):
                if len(imgs) == 1:
                    fname = f"{prefix}.{ext}"
                else:
                    fname = f"{prefix}-{ii + 1:02d}.{ext}"
                path = os.path.join(output_dir, fname)
                with open(path, 'wb') as f:
                    f.write(data)
                files.append(fname)
                print(f"    → {fname} ({len(data):,} bytes)")

        if not files:
            print(f"    (no content)")

        results.append({
            'number': num,
            'title': title,
            'category': entry.get('category', ''),
            'files': files,
        })

    # Manifest
    manifest = {'source': docx_path, 'results': results}
    manifest_path = os.path.join(output_dir, 'manifest.json')
    with open(manifest_path, 'w', encoding='utf-8') as f:
        json.dump(manifest, f, ensure_ascii=False, indent=2)

    total_files = sum(len(r['files']) for r in results)
    print(f"\nDone: {total_files} files from {len(results)} sections → {output_dir}/")


def main():
    parser = argparse.ArgumentParser(description="Extract sections by plan")
    parser.add_argument('docx_path')
    parser.add_argument('--plan', required=True)
    parser.add_argument('--toc', required=True, help="toc.json from Step 1")
    parser.add_argument('--output-dir', '-o', default='data')
    args = parser.parse_args()
    extract_by_plan(args.docx_path, args.plan, args.toc, args.output_dir)


if __name__ == '__main__':
    main()
