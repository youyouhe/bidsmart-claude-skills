"""
Extract neutral/reusable structured data from bid response Word (.docx) files.

Approach: scan ALL tables and paragraphs, match content by keywords,
regardless of table structure (2-col, 4-col, merged cells, etc.).

Only extracts company-level neutral data that can be reused across projects:
  - Company info (name, address, credit code, legal rep, bank, etc.)
  - Personnel (names, education, certifications, experience)
  - Qualifications (ISO certs, CMMI, software copyrights, etc.)
  - Performance/track record (project names, clients, amounts, dates)

Does NOT extract project-specific content:
  - Technical proposals, architecture designs
  - Requirements analysis, implementation plans
  - Project schedules, risk management
  - Pricing (project-specific)

Usage:
    python extract_text.py <docx_path> [--output-dir data] [--index index.json]
"""
import argparse
import json
import os
import re
from datetime import datetime

from docx import Document
from docx.oxml.ns import qn


# ── Keyword → field mapping ──────────────────────────────────────────────

# Keys we want to find in any table cell; value = target field in output JSON.
# Matching is "keyword in cell_text" (substring), so order matters for ambiguity.

COMPANY_KEYS = {
    '供应商名称': 'name',
    '企业名称': 'name',
    '公司名称': 'name',
    '统一社会信用代码': 'credit_code',
    '社会信用代码': 'credit_code',
    '税号': 'credit_code',
    '注册地址': 'registered_address',
    '企业地址': 'registered_address',
    '公司地址': 'registered_address',
    '办公地址': 'office_address',
    '注册资本': 'registered_capital',
    '注册资金': 'registered_capital',
    '注册资本金': 'registered_capital',
    '企业性质': 'company_type',
    '公司类型': 'company_type',
    '注册时间': 'established_date',
    '成立日期': 'established_date',
    '成立时间': 'established_date',
    '邮政编码': 'postal_code',
    '邮编': 'postal_code',
    '电子邮箱': 'email',
    '邮箱': 'email',
    '传真': 'fax',
    '经营范围': 'business_scope',
    '信用等级': 'credit_rating',
    '总人数': 'total_employees',
}

BANK_KEYS = {
    '开户银行': 'bank_name',
    '开户行': 'bank_name',
    '银行账号': 'account_number',
    '账号': 'account_number',
    '开户行联行号': 'bank_code',
    '联行号': 'bank_code',
}

LEGAL_REP_KEYS = {
    '法定代表人': 'name',
    '法人代表': 'name',
}

PERSON_KEYS = {
    '姓名': 'name',
    '姓\t名': 'name',
    '性别': 'gender',
    '年龄': 'age',
    '职务': 'title',
    '职称': 'professional_title',
    '学历': 'education',
    '身份证号码': 'id_number',
    '身份证号': 'id_number',
    '联系电话': 'phone',
    '联系手机': 'phone',
    '电话': 'phone',
    '手机': 'phone',
    '工作年限': 'work_years',
    '工作经验': 'work_experience',
    '执业资格证书': 'certifications',
    '资质证书': 'certifications',
}

# Header keywords that identify a team member list table
TEAM_TABLE_HEADERS = {'人员姓名', '姓名', '学历', '资质证书', '工作经验', '岗位'}

# Header keywords that identify a performance/track record table
PERF_TABLE_HEADERS = {'项目名称', '采购单位', '合同金额', '实施时间', '甲方'}

# ── Helpers ───────────────────────────────────────────────────────────────

PLACEHOLDER_RE = re.compile(r'【.*?】')


def clean(value: str):
    """Return cleaned value, or None if empty/placeholder/meaningless."""
    if not value:
        return None
    v = value.strip().replace('\n', ' ').replace('\t', ' ')
    v = re.sub(r'\s+', ' ', v)
    if not v or v in ('/', '——', '—', '-', '', '无', 'N/A'):
        return None
    if PLACEHOLDER_RE.search(v):
        return None
    return v


def is_noise_value(val: str) -> bool:
    """Check if a value is actually a header/label, not real data."""
    if not val:
        return True
    noise_words = ('复印件', '扫描件', '此处插入', '正面', '反面', '人像面',
                   '国徽面', '（%）', '备注', '说明', '要求', '序号')
    return any(w in val for w in noise_words)


def looks_like_person_name(val: str) -> bool:
    """Rough check: is this a plausible Chinese person name?"""
    if not val:
        return False
    # Reject known table-header words and labels
    NON_NAME_WORDS = {
        '职称', '学历', '时间', '日期', '序号', '编号', '姓名', '性别', '年龄',
        '岗位', '角色', '职务', '备注', '说明', '证书', '资质', '经验', '专业',
        '电话', '手机', '地址', '邮箱', '名称', '金额', '单位', '项目',
    }
    if val in NON_NAME_WORDS:
        return False
    # Reject if contains common label/descriptor fragments
    if any(w in val for w in ('联系', '电话', '手机', '代表', '授权', '委托',
                              '负责', '地址', '邮箱', '编号', '复印件', '扫描件',
                              '证书', '资质', '经验', '学历', '职称')):
        return False
    # Reject date-like patterns: 2025.1至今, 2024.6-2024.12, etc.
    if re.match(r'\d{4}[\.\-/年]', val):
        return False
    # 2-4 Chinese characters, possibly with middot (typical Chinese name)
    if re.fullmatch(r'[\u4e00-\u9fff·]{2,6}', val):
        return True
    # Short alphanumeric (foreign names) — must not contain common noise
    if len(val) <= 20 and re.search(r'[a-zA-Z\u4e00-\u9fff]', val):
        return True
    return False


def dedup_row(row):
    """Deduplicate cells from merged columns (python-docx repeats them)."""
    seen = set()
    cells = []
    for cell in row.cells:
        cid = id(cell._tc)
        if cid not in seen:
            seen.add(cid)
            cells.append(cell.text.strip())
    return cells


def get_contexts(doc):
    """Map table_index → last few paragraphs before that table."""
    contexts = {}
    ti = 0
    recent = []
    for child in doc.element.body:
        if child.tag == qn('w:p'):
            texts = []
            for r in child.iter(qn('w:t')):
                if r.text:
                    texts.append(r.text)
            t = ''.join(texts).strip()
            if t:
                recent.append(t)
                if len(recent) > 5:
                    recent = recent[-5:]
        elif child.tag == qn('w:tbl'):
            contexts[ti] = list(recent)
            ti += 1
    return contexts


# ── KV pair scanner ───────────────────────────────────────────────────────

def scan_kv_pairs(table):
    """
    Yield (key_text, value_text) from ANY table layout.

    Handles:
    - 2-col: key | value
    - 4-col: key1 | val1 | key2 | val2
    - N-col with merged cells
    - Person cards with interleaved k-v in a single row
    """
    for row in table.rows:
        cells = dedup_row(row)
        if len(cells) < 2:
            continue
        # Try adjacent pairs: cells[0]→cells[1], cells[2]→cells[3], ...
        i = 0
        while i < len(cells) - 1:
            key = cells[i].strip()
            val = cells[i + 1].strip()
            if key and len(key) < 30:  # keys are short labels
                yield key, val
            i += 2
        # Also try cells[1]→cells[2] for offset patterns
        i = 1
        while i < len(cells) - 1:
            key = cells[i].strip()
            val = cells[i + 1].strip()
            if key and len(key) < 30 and len(key) > 1:
                yield key, val
            i += 2


def scan_header_rows(table):
    """
    If table has a recognizable header row, yield (header_list, data_rows).
    """
    rows = []
    for row in table.rows:
        rows.append(dedup_row(row))
    if not rows or len(rows) < 2:
        return None, None
    return rows[0], rows[1:]


# ── Extraction functions ─────────────────────────────────────────────────

def extract_company(doc, contexts):
    """Scan entire doc for company-level info."""
    company = {
        'name': None, 'credit_code': None, 'registered_address': None,
        'office_address': None, 'registered_capital': None,
        'company_type': None, 'established_date': None,
        'postal_code': None, 'email': None, 'phone': None, 'fax': None,
        'business_scope': None, 'credit_rating': None, 'total_employees': None,
        'legal_representative': None,
        'bank': {
            'bank_name': None, 'account_number': None, 'bank_code': None,
        },
    }

    for ti, table in enumerate(doc.tables):
        for key, val in scan_kv_pairs(table):
            cv = clean(val)
            if not cv or is_noise_value(cv):
                continue

            # Company fields
            for kw, field in COMPANY_KEYS.items():
                if kw in key:
                    # For 'name' field, prefer exact '供应商名称' over '公司名称'
                    # to avoid matching '母公司及控股公司名称' etc.
                    if field == 'name':
                        if key.startswith('母公司') or '控股' in key:
                            break  # skip this match
                    if company.get(field) is None:
                        company[field] = cv
                    break

            # Bank fields
            for kw, field in BANK_KEYS.items():
                if kw in key:
                    if company['bank'].get(field) is None:
                        company['bank'][field] = cv
                    break

            # Legal representative: must look like a name (short, no noise words)
            for kw, field in LEGAL_REP_KEYS.items():
                if kw in key and company['legal_representative'] is None:
                    if looks_like_person_name(cv):
                        company['legal_representative'] = cv
                    break

            # Phone: only capture if no phone yet
            if company['phone'] is None:
                if key in ('联系电话', '联系手机', '电话', '手机', '联系手机'):
                    company['phone'] = cv

    return company


def extract_personnel(doc, contexts):
    """Extract personnel from person-card tables and team summary tables."""
    personnel = []
    seen_names = set()

    for ti, table in enumerate(doc.tables):
        ctx_lines = contexts.get(ti, [])
        ctx_text = ' '.join(ctx_lines)

        # ── Try header+rows pattern (team summary table) ──
        header, data_rows = scan_header_rows(table)
        if header and data_rows:
            header_text = ' '.join(header)
            # Check if this looks like a team member list
            matches = sum(1 for h in TEAM_TABLE_HEADERS if any(h in c for c in header))
            if matches >= 2 and ('人员姓名' in header_text or '姓名' in header_text):
                # Build column mapping from header
                col_map = {}
                for ci, h in enumerate(header):
                    if '姓名' in h or '人员姓名' in h:
                        col_map['name'] = ci
                    elif '学历' in h:
                        col_map['education'] = ci
                    elif '资质' in h or '证书' in h:
                        col_map['certifications'] = ci
                    elif '工作经验' in h or '经验' in h:
                        col_map['experience'] = ci
                    elif '岗位' in h or '角色' in h:
                        col_map['role'] = ci

                if 'name' in col_map:
                    for row in data_rows:
                        name = clean(row[col_map['name']]) if col_map['name'] < len(row) else None
                        if not name or not looks_like_person_name(name):
                            continue
                        if name in seen_names:
                            continue
                        seen_names.add(name)
                        person = {
                            'name': name,
                            'role': clean(row[col_map['role']]) if 'role' in col_map and col_map['role'] < len(row) else None,
                            'education': clean(row[col_map['education']]) if 'education' in col_map and col_map['education'] < len(row) else None,
                            'certifications': clean(row[col_map['certifications']]) if 'certifications' in col_map and col_map['certifications'] < len(row) else None,
                            'experience': clean(row[col_map['experience']]) if 'experience' in col_map and col_map['experience'] < len(row) else None,
                        }
                        personnel.append(person)
                    continue  # done with this table

        # ── Try KV pattern (person card) ──
        person = {}
        for key, val in scan_kv_pairs(table):
            cv = clean(val)
            if not cv:
                continue
            for kw, field in PERSON_KEYS.items():
                if kw in key:
                    if field not in person:
                        person[field] = cv
                    break

        # Only keep if we found a name that looks like a real person name
        if person.get('name') and person['name'] not in seen_names and looks_like_person_name(person['name']):
            # Determine role from context
            role = _guess_role(ctx_lines)
            if role:
                person['role'] = role
            seen_names.add(person['name'])

            # Also check for project experience rows in the same table
            experience_items = []
            for row in table.rows:
                cells = dedup_row(row)
                # Look for rows with date-like first cell and project descriptions
                if cells and re.match(r'\d{4}', cells[0].strip()):
                    proj_text = ' '.join(c for c in cells[1:] if clean(c))
                    if proj_text:
                        experience_items.append(f"{cells[0].strip()}: {proj_text[:100]}")

            if experience_items:
                person['project_experience'] = experience_items

            personnel.append(person)

    return personnel


def _guess_role(ctx_lines):
    """Guess person's role from context paragraphs (last line wins)."""
    role_keywords = [
        '项目负责人', '项目经理', '系统架构师', '架构师',
        '需求分析师', '需求分析', '测试工程师', '测试',
        '实施工程师', '实施', '运维工程师', '运维',
        '开发工程师', '开发', '前端', '后端',
    ]
    for line in reversed(ctx_lines):
        for rk in role_keywords:
            if rk in line:
                return rk
    return None


def extract_performance(doc, contexts):
    """Extract project track record / performance."""
    projects = []

    for ti, table in enumerate(doc.tables):
        header, data_rows = scan_header_rows(table)
        if not header or not data_rows:
            continue

        header_text = ' '.join(header)
        matches = sum(1 for h in PERF_TABLE_HEADERS if any(h in c for c in header))
        if matches < 2:
            continue

        # Build column mapping
        col_map = {}
        for ci, h in enumerate(header):
            if '项目名称' in h:
                col_map['project_name'] = ci
            elif '采购单位' in h or '采购人' in h or '甲方' in h or '用户' in h:
                col_map['client'] = ci
            elif '合同金额' in h or '金额' in h:
                col_map['amount'] = ci
            elif '时间' in h or '日期' in h:
                col_map['date'] = ci

        for row in data_rows:
            proj = {}
            for field, ci in col_map.items():
                if ci < len(row):
                    proj[field] = clean(row[ci])
            # Skip empty rows
            if proj.get('project_name'):
                projects.append(proj)

    return projects


def extract_qualifications(doc):
    """Scan paragraphs and tables for qualification/certification mentions."""
    certs = []
    found = set()

    cert_keywords = {
        'ISO9001': '质量管理体系', 'ISO 9001': '质量管理体系',
        'ISO14001': '环境管理体系', 'ISO 14001': '环境管理体系',
        'ISO27001': '信息安全管理体系', 'ISO 27001': '信息安全管理体系',
        'ISO20000': '信息技术服务管理体系', 'ISO 20000': '信息技术服务管理体系',
        'ISO45001': '职业健康安全管理体系', 'ISO 45001': '职业健康安全管理体系',
        'OHSAS18001': '职业健康安全管理体系',
        'CMMI': '能力成熟度模型',
        '高新技术企业': '高新技术企业认定',
        '软件著作权': '软件著作权',
        'ITSS': '信息技术服务标准',
    }

    # Scan all text (paragraphs + table cells)
    all_text = []
    for p in doc.paragraphs:
        if p.text.strip():
            all_text.append(p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    all_text.append(cell.text.strip())

    for text in all_text:
        for kw, desc in cert_keywords.items():
            if kw in text and kw not in found:
                found.add(kw)
                certs.append({'name': kw, 'description': desc})

    return certs


# ── Main ──────────────────────────────────────────────────────────────────

def extract_all(docx_path: str, output_dir: str, index_path: str = None):
    os.makedirs(output_dir, exist_ok=True)

    print(f"Loading {docx_path}...")
    doc = Document(docx_path)
    print(f"  {len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables")

    contexts = get_contexts(doc)
    source = os.path.basename(docx_path)
    ts = datetime.now().isoformat()

    # ── Company ──
    print("\nExtracting company info...")
    company = extract_company(doc, contexts)
    company['source_file'] = source
    company['extracted_at'] = ts
    _save(os.path.join(output_dir, 'company.json'), company)
    print(f"  名称: {company.get('name') or '—'}")
    print(f"  信用代码: {company.get('credit_code') or '—'}")
    print(f"  法人: {company.get('legal_representative') or '—'}")
    print(f"  地址: {company.get('registered_address') or '—'}")
    print(f"  银行: {company['bank'].get('bank_name') or '—'}")

    # ── Personnel ──
    print("\nExtracting personnel...")
    personnel = extract_personnel(doc, contexts)
    _save(os.path.join(output_dir, 'personnel.json'), {
        'personnel': personnel, 'source_file': source, 'extracted_at': ts,
    })
    print(f"  {len(personnel)} people found")
    for p in personnel:
        role = p.get('role') or ''
        certs = p.get('certifications') or p.get('professional_title') or ''
        print(f"    {p['name']} | {role} | {certs}"[:80])

    # ── Performance ──
    print("\nExtracting performance...")
    projects = extract_performance(doc, contexts)
    _save(os.path.join(output_dir, 'performance.json'), {
        'projects': projects, 'source_file': source, 'extracted_at': ts,
    })
    print(f"  {len(projects)} projects found")
    for proj in projects:
        print(f"    {proj.get('project_name', '—')[:50]} | {proj.get('client', '—')[:20]}")

    # ── Qualifications ──
    print("\nExtracting qualifications...")
    quals = extract_qualifications(doc)
    _save(os.path.join(output_dir, 'qualifications.json'), {
        'certifications': quals, 'source_file': source, 'extracted_at': ts,
    })
    print(f"  {len(quals)} certifications found")
    for q in quals:
        print(f"    {q['name']} ({q['description']})")

    # ── Cross-reference with image index ──
    if index_path and os.path.exists(index_path):
        print(f"\nCross-referencing with {index_path}...")
        _cross_ref(output_dir, index_path)

    # ── Report ──
    _report(output_dir, company, personnel, projects, quals, source)
    print(f"\nDone. Output: {output_dir}/")


def _save(path, data):
    with open(path, 'w', encoding='utf-8') as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def _cross_ref(output_dir, index_path):
    """Link qualifications and performance to image index entries."""
    with open(index_path, 'r', encoding='utf-8') as f:
        index = json.load(f)
    docs = index.get('documents', [])

    # quals
    qpath = os.path.join(output_dir, 'qualifications.json')
    if os.path.exists(qpath):
        with open(qpath, 'r', encoding='utf-8') as f:
            qdata = json.load(f)
        for cert in qdata.get('certifications', []):
            name_lower = cert['name'].lower().replace(' ', '')
            for d in docs:
                tags = ' '.join(d.get('searchable_tags', [])).lower()
                dtype = d.get('type', '').lower()
                if name_lower in tags.replace(' ', '') or name_lower in dtype.replace(' ', ''):
                    cert['image_ref'] = d['id']
                    break
        _save(qpath, qdata)

    # performance
    ppath = os.path.join(output_dir, 'performance.json')
    if os.path.exists(ppath):
        with open(ppath, 'r', encoding='utf-8') as f:
            pdata = json.load(f)
        contract_docs = [d for d in docs if d['category'] == '业绩证明']
        for pi, proj in enumerate(pdata.get('projects', [])):
            if pi < len(contract_docs):
                proj['image_ref'] = contract_docs[pi]['id']
        _save(ppath, pdata)


def _report(output_dir, company, personnel, projects, quals, source):
    lines = [
        '# 资料提取报告', '',
        f'- 来源: `{source}`',
        f'- 时间: {datetime.now().strftime("%Y-%m-%d %H:%M")}', '',
        '## 公司信息', '',
    ]
    for k, v in company.items():
        if k in ('source_file', 'extracted_at', 'bank'):
            continue
        if isinstance(v, str) and v:
            lines.append(f'- {k}: {v}')
    bank = company.get('bank', {})
    if any(bank.values()):
        lines.append(f'- 开户行: {bank.get("bank_name") or "—"}')
        lines.append(f'- 账号: {bank.get("account_number") or "—"}')

    if personnel:
        lines.extend(['', '## 人员信息', '',
                       '| 姓名 | 角色 | 学历 | 资质 |',
                       '|------|------|------|------|'])
        for p in personnel:
            name = p.get('name', '—')
            role = p.get('role', '—') or '—'
            edu = p.get('education', '—') or '—'
            cert = p.get('certifications') or p.get('professional_title') or '—'
            lines.append(f'| {name} | {role} | {edu} | {cert} |')

    if projects:
        lines.extend(['', '## 业绩信息', '',
                       '| 项目 | 客户 | 金额 | 时间 |',
                       '|------|------|------|------|'])
        for proj in projects:
            lines.append(f'| {proj.get("project_name", "—")} | {proj.get("client", "—")} | {proj.get("amount", "—")} | {proj.get("date", "—")} |')

    if quals:
        lines.extend(['', '## 资质证书', ''])
        for q in quals:
            lines.append(f'- {q["name"]} ({q["description"]})')

    with open(os.path.join(output_dir, 'extraction_report.md'), 'w', encoding='utf-8') as f:
        f.write('\n'.join(lines))


if __name__ == '__main__':
    parser = argparse.ArgumentParser()
    parser.add_argument('docx_path')
    parser.add_argument('--output-dir', default='data')
    parser.add_argument('--index', default=None)
    args = parser.parse_args()
    extract_all(args.docx_path, args.output_dir, args.index)
