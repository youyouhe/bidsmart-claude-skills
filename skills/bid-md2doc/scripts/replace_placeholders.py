#!/usr/bin/env python3
"""
占位符替换工具 - 根据Excel清单替换Word文档中的占位符为本地图片

功能：
1. 读取 占位符清单.xlsx
2. 根据用户填写的本地图片路径
3. 在Word文档中查找对应占位符并替换为图片
4. 生成新的Word文档（原文件不变）
5. 完全离线工作，不依赖任何API

使用方法：
    python replace_placeholders.py

前置条件：
    - 当前目录下有 占位符清单.xlsx
    - Excel中已填写本地图片路径
    - 当前目录下有原始Word文档

输出：
    {原文件名}-已替换.docx
"""

import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Inches, Cm
from docx.oxml.ns import qn
from openpyxl import load_workbook
from PIL import Image


def load_placeholder_mapping(excel_path: str) -> dict:
    """
    从Excel清单加载占位符与本地图片路径的映射

    Args:
        excel_path: Excel文件路径

    Returns:
        字典 {占位符原文: 本地图片路径}
    """
    try:
        wb = load_workbook(excel_path)
        ws = wb.active
    except Exception as e:
        print(f"❌ 无法打开Excel文件: {e}")
        sys.exit(1)

    mapping = {}
    skipped = []

    print("📊 正在读取占位符清单...")

    # 跳过标题行，从第2行开始读取
    for row_num, row in enumerate(ws.iter_rows(min_row=2, values_only=True), 2):
        if len(row) < 4:
            continue

        index, placeholder, material_name, image_path = row[0], row[1], row[2], row[3]

        if not placeholder:
            continue

        # 如果用户没有填写图片路径，跳过
        if not image_path or str(image_path).strip() == '':
            skipped.append((index, placeholder))
            continue

        image_path = str(image_path).strip()

        # 检查图片文件是否存在
        if not Path(image_path).exists():
            print(f"⚠️  行 {row_num}: 图片文件不存在: {image_path}")
            skipped.append((index, placeholder))
            continue

        # 检查文件格式
        if not image_path.lower().endswith(('.png', '.jpg', '.jpeg')):
            print(f"⚠️  行 {row_num}: 不支持的图片格式: {image_path}")
            skipped.append((index, placeholder))
            continue

        mapping[placeholder] = image_path
        print(f"✓ 占位符 #{index}: {placeholder} → {Path(image_path).name}")

    print(f"\n✅ 成功加载 {len(mapping)} 个映射")
    if skipped:
        print(f"ℹ️  跳过 {len(skipped)} 个未填写或无效的占位符")

    return mapping


def get_image_size(image_path: str, max_width_cm: float = 15.0) -> tuple:
    """
    计算图片插入Word时的合适尺寸

    Args:
        image_path: 图片路径
        max_width_cm: 最大宽度（厘米）

    Returns:
        (宽度, 高度) 单位为Cm
    """
    try:
        with Image.open(image_path) as img:
            width_px, height_px = img.size
    except Exception as e:
        print(f"⚠️  无法读取图片尺寸: {e}")
        # 返回默认尺寸
        return (Cm(10), Cm(7))

    # 计算宽高比
    aspect_ratio = height_px / width_px

    # 限制最大宽度
    width_cm = min(max_width_cm, 15.0)
    height_cm = width_cm * aspect_ratio

    # 限制最大高度（A4纸可用高度约25cm）
    if height_cm > 20.0:
        height_cm = 20.0
        width_cm = height_cm / aspect_ratio

    return (Cm(width_cm), Cm(height_cm))


def replace_placeholders_in_docx(
    docx_path: str,
    mapping: dict,
    output_path: str
) -> dict:
    """
    在Word文档中替换占位符为图片

    Args:
        docx_path: 原始Word文档路径
        mapping: 占位符与图片路径的映射
        output_path: 输出Word文档路径

    Returns:
        统计信息 {replaced_count, failed_count, not_found_count}
    """
    try:
        doc = Document(docx_path)
    except Exception as e:
        print(f"❌ 无法打开Word文档: {e}")
        sys.exit(1)

    stats = {
        'replaced_count': 0,
        'failed_count': 0,
        'not_found_count': 0,
    }

    replaced_placeholders = set()

    print(f"\n📄 正在处理文档: {Path(docx_path).name}")
    print("=" * 60)

    # 遍历所有段落
    for para_num, para in enumerate(doc.paragraphs, 1):
        text = para.text

        # 查找当前段落中的所有占位符
        for placeholder in mapping.keys():
            if placeholder in text:
                image_path = mapping[placeholder]

                try:
                    # 清空段落文本
                    para.clear()

                    # 获取图片尺寸
                    width, height = get_image_size(image_path)

                    # 插入图片
                    run = para.add_run()
                    run.add_picture(image_path, width=width, height=height)

                    # 居中对齐
                    para.alignment = 1  # 1 = CENTER

                    stats['replaced_count'] += 1
                    replaced_placeholders.add(placeholder)

                    print(f"✓ 替换成功 [段落 {para_num}]: {placeholder}")
                    print(f"  → 图片: {Path(image_path).name} ({width.cm:.1f}cm × {height.cm:.1f}cm)")

                except Exception as e:
                    stats['failed_count'] += 1
                    print(f"✗ 替换失败 [段落 {para_num}]: {placeholder}")
                    print(f"  → 错误: {e}")

    # 统计未找到的占位符
    not_found = set(mapping.keys()) - replaced_placeholders
    stats['not_found_count'] = len(not_found)

    if not_found:
        print("\n⚠️  以下占位符在文档中未找到：")
        for ph in sorted(not_found):
            print(f"  - {ph}")

    # 保存新文档
    try:
        doc.save(output_path)
        print(f"\n✅ 已保存替换后的文档: {output_path}")
    except Exception as e:
        print(f"\n❌ 保存文档失败: {e}")
        sys.exit(1)

    return stats


def find_docx_file(directory: Path) -> Path:
    """
    在目录中查找Word文档（排除已替换的文件）

    Args:
        directory: 目录路径

    Returns:
        找到的Word文档路径
    """
    docx_files = [
        f for f in directory.glob("*.docx")
        if not f.name.startswith("~$") and "已替换" not in f.name
    ]

    if not docx_files:
        print("❌ 当前目录下没有找到Word文档")
        sys.exit(1)

    if len(docx_files) == 1:
        return docx_files[0]

    # 多个文档，让用户选择
    print("\n找到多个Word文档，请选择：")
    for i, f in enumerate(docx_files, 1):
        print(f"{i}. {f.name}")

    try:
        choice = int(input("\n请输入序号 (1-{}): ".format(len(docx_files))))
        if 1 <= choice <= len(docx_files):
            return docx_files[choice - 1]
        else:
            print("❌ 无效的选择")
            sys.exit(1)
    except (ValueError, KeyboardInterrupt):
        print("\n❌ 已取消")
        sys.exit(1)


def main():
    print("\n" + "=" * 60)
    print("🔄 占位符替换工具 v1.0")
    print("=" * 60 + "\n")

    # 当前目录
    current_dir = Path.cwd()

    # 查找Excel清单
    excel_path = current_dir / "占位符清单.xlsx"
    if not excel_path.exists():
        print("❌ 未找到 占位符清单.xlsx")
        print("请先运行 extract_placeholders.py 生成清单")
        sys.exit(1)

    # 加载映射关系
    mapping = load_placeholder_mapping(str(excel_path))

    if not mapping:
        print("\n❌ Excel中没有有效的占位符映射")
        print("请在"本地图片路径"列填写图片文件路径")
        sys.exit(1)

    # 查找Word文档
    docx_path = find_docx_file(current_dir)
    print(f"\n📄 使用文档: {docx_path.name}\n")

    # 生成输出文件名
    output_path = docx_path.parent / f"{docx_path.stem}-已替换.docx"

    # 执行替换
    stats = replace_placeholders_in_docx(str(docx_path), mapping, str(output_path))

    # 输出统计
    print("\n" + "=" * 60)
    print("📊 替换统计")
    print("=" * 60)
    print(f"✅ 成功替换: {stats['replaced_count']} 个")
    print(f"❌ 替换失败: {stats['failed_count']} 个")
    print(f"⚠️  未找到: {stats['not_found_count']} 个")
    print("=" * 60)

    if stats['replaced_count'] > 0:
        print(f"\n✅ 替换完成！请查看: {output_path.name}")
    else:
        print("\n⚠️  没有成功替换任何占位符")

    print()


if __name__ == '__main__':
    main()
